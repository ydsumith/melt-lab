from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn

OUT = Path(r"C:\Users\syesudasan\OneDrive - University of New Haven\RESEARCH\2. Journals\in preparation\0. Review paper Sumith\manuscript\submission")
OUT.mkdir(parents=True, exist_ok=True)
path = OUT / "Cover_Letter_Seminars_in_Thrombosis_and_Hemostasis.docx"

doc = Document()
section = doc.sections[0]
section.page_width = Inches(8.5)
section.page_height = Inches(11)
section.top_margin = Inches(0.65)
section.bottom_margin = Inches(0.65)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)
section.header_distance = Inches(0.49)
section.footer_distance = Inches(0.49)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Calibri"
normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
normal.font.size = Pt(10.5)
normal.paragraph_format.space_before = Pt(0)
normal.paragraph_format.space_after = Pt(6)
normal.paragraph_format.line_spacing = 1.05

def set_font(run, size=11, bold=False, color="000000", italic=False):
    run.font.name = "Calibri"
    run._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    run._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)

def para(text="", after=6, before=0, align=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(before)
    p.paragraph_format.space_after = Pt(after)
    p.paragraph_format.line_spacing = 1.05
    if align is not None:
        p.alignment = align
    r = p.add_run(text)
    set_font(r)
    return p

# Traditional academic letterhead (restrained memo-masthead adaptation).
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(2)
r = p.add_run("SUMITH YESUDASAN")
set_font(r, size=16, bold=True, color="1F4D78")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("Department of Mechanical and Industrial Engineering")
set_font(r, size=10.5, bold=True)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("University of New Haven | 300 Boston Post Road | West Haven, CT 06516, USA")
set_font(r, size=10, color="555555")

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(12)
r = p.add_run("syesudasan@newhaven.edu")
set_font(r, size=10, color="555555")

para("July 19, 2026", after=12)
para("Emmanuel J. Favaloro, PhD, FFSc (RCPA)", after=1)
para("Editor-in-Chief", after=1)
para("Seminars in Thrombosis and Hemostasis", after=1)
para("Institute of Clinical Pathology and Medical Research", after=1)
para("Westmead Hospital", after=1)
para("Westmead, New South Wales, Australia", after=12)

para("Dear Professor Favaloro:", after=8)

title = "Multiscale Computational Modeling of Fibrin Polymerization: Molecular Mechanisms, Network Assembly, and Emerging Challenges"
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.05
r = p.add_run("Re: Submission of the Review Article, ")
set_font(r, bold=True)
r = p.add_run(f'“{title}”')
set_font(r, bold=True, italic=True)

para(
    "Please consider the enclosed manuscript for publication as a Review Article in Seminars in Thrombosis and Hemostasis. "
    "The article critically evaluates computational approaches to fibrin formation across atomistic, coarse-grained, mesoscopic, fiber-network, and continuum scales."
)

para(
    "The review is organized around a central question: when can a multiscale fibrin model be considered predictive rather than merely representational? "
    "It examines parameter provenance, thermodynamic and kinetic consistency, hierarchical validation, uncertainty transfer, and the treatment of pathological perturbations. "
    "The manuscript also distinguishes thrombin-mediated activation, reversible knob-hole recognition, lateral association, factor XIIIa-mediated crosslinking, mechanical damage, and fibrinolytic degradation as separate mechanisms."
)

para(
    "This synthesis should be relevant to the journal’s readership because it connects molecular mechanisms of fibrin polymerization with clot architecture, mechanics, permeability, fibrinolysis, and failure, while identifying practical priorities for experimentally constrained computational modeling."
)

para(
    "The manuscript is original, has not been published previously, and is not under consideration by another journal. I am the sole author, have approved the submitted version, and declare no competing interests. No new experimental data were generated for this review."
)

para(
    "Thank you for considering this manuscript. I would be grateful for the opportunity to have it evaluated for publication in Seminars in Thrombosis and Hemostasis.",
    after=12,
)

para("Sincerely,", after=14)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(1)
r = p.add_run("Sumith Yesudasan")
set_font(r, bold=True)
para("Department of Mechanical and Industrial Engineering", after=1)
para("University of New Haven", after=1)
para("syesudasan@newhaven.edu", after=0)

# Quiet footer.
footer = section.footer
fp = footer.paragraphs[0]
fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = fp.add_run("Cover Letter | Seminars in Thrombosis and Hemostasis")
set_font(r, size=8.5, color="777777")

doc.core_properties.title = "Cover Letter - Seminars in Thrombosis and Hemostasis"
doc.core_properties.author = "Sumith Yesudasan"
doc.core_properties.subject = title
doc.save(path)
print(path)
